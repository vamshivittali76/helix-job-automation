"""
LLM-powered resume tailoring with strict ATS compliance.

Takes the user's base resume text, a job description, and profile data,
then produces a customized ATS-friendly resume as .docx (and optionally .pdf).

RULES (from PROJECT_CONTEXT.xml):
  - NEVER fabricate skills, titles, companies, or achievements
  - Only rephrase, reorder, and emphasize existing experience
  - Mirror exact keywords from the job description
  - Use standard section headings, simple formatting, single column
"""

import json
from pathlib import Path
from typing import Any, Optional

OUTPUT_DIR = Path(__file__).parent.parent.parent / "output" / "resumes"
TEMPLATE_DIR = Path(__file__).parent.parent.parent / "templates"


def extract_resume_text(docx_path: str | Path) -> str:
    """Extract plain text from a .docx resume."""
    from docx import Document
    doc = Document(str(docx_path))
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def tailor_resume(
    client: Any,
    base_resume_text: str,
    job_title: str,
    job_company: str,
    job_description: str,
    profile: dict,
    company_research: str = "",
    model: str = "gpt-4o-mini",
) -> dict:
    """
    Generate tailored resume content via LLM.

    Returns dict with keys: professional_summary, skills, experience, education,
    certifications, projects, keyword_additions
    """
    research_ctx = ""
    if company_research:
        try:
            r = json.loads(company_research)
            research_ctx = f"\nCompany info: {r.get('summary', '')} Industry: {r.get('industry', '')}"
        except (json.JSONDecodeError, TypeError):
            pass

    prompt = (
        "You are an expert resume writer. Tailor this resume for the target job.\n\n"
        "STRICT RULES:\n"
        "- NEVER invent new skills, job titles, companies, or achievements\n"
        "- Only rephrase, reorder, and emphasize what already exists in the base resume\n"
        "- Mirror exact keywords from the job description where truthful\n"
        "- Quantify achievements with numbers where they already exist\n"
        "- Professional Summary must be 2-3 lines tailored to this specific role\n\n"
        f"TARGET JOB:\nTitle: {job_title}\nCompany: {job_company}\n"
        f"Description: {job_description[:2500]}\n{research_ctx}\n\n"
        f"BASE RESUME:\n{base_resume_text[:4000]}\n\n"
        "Respond with valid JSON:\n"
        "{\n"
        '  "professional_summary": "<2-3 tailored sentences>",\n'
        '  "skills": ["skill1", "skill2", ...],\n'
        '  "experience": [\n'
        '    {"title": "...", "company": "...", "dates": "...", '
        '"bullets": ["achievement1", "achievement2", ...]},\n'
        "    ...\n"
        "  ],\n"
        '  "education": [{"degree": "...", "school": "...", "year": "..."}],\n'
        '  "certifications": ["cert1", ...],\n'
        '  "projects": [{"name": "...", "description": "...", "tech": "..."}]\n'
        "}"
    )

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an ATS-optimized resume writer. You NEVER fabricate experience. "
                        "You only rephrase and reorder existing content to match the target role. "
                        "Return only valid JSON."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            max_tokens=2000,
        )

        text = response.choices[0].message.content.strip()
        if text.startswith("```"):
            text = text.split("\n", 1)[1] if "\n" in text else text[3:]
            text = text.rsplit("```", 1)[0]

        return json.loads(text)
    except Exception as e:
        return {"error": str(e)}


def build_resume_docx(
    content: dict,
    profile: dict,
    output_path: str | Path,
) -> Path:
    """Build an ATS-compliant .docx resume from tailored content."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    personal = profile.get("personal", {})
    name = f"{personal.get('first_name', '')} {personal.get('last_name', '')}".strip()
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.font.size = Pt(16)
        run.font.bold = True

    contact_parts = []
    if personal.get("email"):
        contact_parts.append(personal["email"])
    if personal.get("phone"):
        contact_parts.append(personal["phone"])
    if personal.get("location"):
        contact_parts.append(personal["location"])
    if personal.get("linkedin"):
        contact_parts.append(personal["linkedin"])
    if contact_parts:
        p = doc.add_paragraph(" | ".join(contact_parts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style.font.size = Pt(10)

    def add_section_heading(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        p.space_after = Pt(4)

    if content.get("professional_summary"):
        add_section_heading("Professional Summary")
        doc.add_paragraph(content["professional_summary"])

    if content.get("experience"):
        add_section_heading("Experience")
        for exp in content["experience"]:
            p = doc.add_paragraph()
            title_run = p.add_run(exp.get("title", ""))
            title_run.bold = True
            p.add_run(f" | {exp.get('company', '')}")
            if exp.get("dates"):
                p.add_run(f" | {exp['dates']}")
            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.text = bullet
                bp.paragraph_format.space_after = Pt(2)

    if content.get("skills"):
        add_section_heading("Skills")
        skills_text = " • ".join(content["skills"])
        doc.add_paragraph(skills_text)

    if content.get("education"):
        add_section_heading("Education")
        for edu in content["education"]:
            parts = [edu.get("degree", "")]
            if edu.get("school"):
                parts.append(edu["school"])
            if edu.get("year"):
                parts.append(str(edu["year"]))
            doc.add_paragraph(" | ".join(p for p in parts if p))

    if content.get("certifications"):
        certs = [c for c in content["certifications"] if c]
        if certs:
            add_section_heading("Certifications")
            for cert in certs:
                doc.add_paragraph(f"• {cert}")

    if content.get("projects"):
        projects = [p for p in content["projects"] if p.get("name")]
        if projects:
            add_section_heading("Projects")
            for proj in projects:
                p = doc.add_paragraph()
                run = p.add_run(proj.get("name", ""))
                run.bold = True
                if proj.get("tech"):
                    p.add_run(f" ({proj['tech']})")
                if proj.get("description"):
                    doc.add_paragraph(proj["description"])

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def convert_to_pdf(docx_path: str | Path) -> Optional[Path]:
    """Convert .docx to .pdf using docx2pdf (requires MS Word on Windows)."""
    try:
        from docx2pdf import convert
        pdf_path = Path(str(docx_path).replace(".docx", ".pdf"))
        convert(str(docx_path), str(pdf_path))
        return pdf_path
    except Exception:
        return None


def generate_resume_for_job(
    client: Any,
    job: dict,
    profile: dict,
    base_resume_path: str | Path = None,
    model: str = "gpt-4o-mini",
) -> Optional[Path]:
    """Full pipeline: read base resume -> tailor -> build docx -> return path."""
    base_path = Path(base_resume_path) if base_resume_path else TEMPLATE_DIR / "base_resume.docx"
    if not base_path.exists():
        return None

    base_text = extract_resume_text(base_path)
    if not base_text.strip():
        return None

    content = tailor_resume(
        client=client,
        base_resume_text=base_text,
        job_title=job.get("title", ""),
        job_company=job.get("company", ""),
        job_description=job.get("description", ""),
        profile=profile,
        company_research=job.get("company_research", ""),
        model=model,
    )

    if content.get("error"):
        return None

    personal = profile.get("personal", {})
    first = personal.get("first_name", "Resume")
    last = personal.get("last_name", "")
    company_clean = (job.get("company") or "Company").replace(" ", "_").replace("/", "_")
    filename = f"{first}_{last}_Resume_{company_clean}.docx"
    output_path = OUTPUT_DIR / filename

    build_resume_docx(content, profile, output_path)
    convert_to_pdf(output_path)

    return output_path
