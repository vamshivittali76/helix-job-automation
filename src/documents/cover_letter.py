"""
LLM-powered cover letter generation with ATS compliance.

Generates personalized, ATS-friendly cover letters that incorporate
company research and mirror job description keywords.
"""

import json
from pathlib import Path
from typing import Any, Optional

OUTPUT_DIR = Path(__file__).parent.parent.parent / "output" / "cover_letters"


def generate_cover_letter_content(
    client: Any,
    job_title: str,
    job_company: str,
    job_description: str,
    profile: dict,
    company_research: str = "",
    model: str = "gpt-4o-mini",
) -> dict:
    """
    Generate tailored cover letter text via LLM.

    Returns dict with: greeting, opening, body, closing, sign_off
    """
    personal = profile.get("personal", {})
    name = f"{personal.get('first_name', '')} {personal.get('last_name', '')}".strip()

    skills_flat = []
    for items in profile.get("skills", {}).values():
        if isinstance(items, list):
            skills_flat.extend(items)

    research_ctx = ""
    if company_research:
        try:
            r = json.loads(company_research) if isinstance(company_research, str) else company_research
            research_ctx = (
                f"\nCompany research: {r.get('summary', '')} "
                f"Culture: {r.get('culture', '')} "
                f"Talking points: {', '.join(r.get('talking_points', []))}"
            )
        except (json.JSONDecodeError, TypeError):
            pass

    prompt = (
        "Write a concise, professional cover letter for this job application.\n\n"
        "RULES:\n"
        "- 3 paragraphs max (opening, body with evidence, closing)\n"
        "- Mirror exact keywords from the job description naturally\n"
        "- Reference specific company details from research if available\n"
        "- NEVER fabricate experience or skills not in the profile\n"
        "- Use 'Dear Hiring Team' unless a specific name is known\n"
        "- Keep it under 350 words total\n"
        "- Be genuine, not generic\n\n"
        f"JOB:\nTitle: {job_title}\nCompany: {job_company}\n"
        f"Description: {job_description[:2000]}\n{research_ctx}\n\n"
        f"CANDIDATE:\nName: {name}\n"
        f"Key Skills: {', '.join(skills_flat[:20])}\n"
        f"Experience: {profile.get('years_of_experience', '')} years, "
        f"{profile.get('experience_level', 'mid')} level\n\n"
        "Respond with valid JSON:\n"
        "{\n"
        '  "greeting": "Dear Hiring Team,",\n'
        '  "opening": "<paragraph 1: hook + why this role at this company>",\n'
        '  "body": "<paragraph 2: evidence of skills, achievements matching JD>",\n'
        '  "closing": "<paragraph 3: enthusiasm + call to action>",\n'
        '  "sign_off": "Sincerely,"\n'
        "}"
    )

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a career counselor writing ATS-friendly cover letters. "
                        "Be concise, genuine, and keyword-rich. Return only valid JSON."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.4,
            max_tokens=800,
        )

        text = response.choices[0].message.content.strip()
        if text.startswith("```"):
            text = text.split("\n", 1)[1] if "\n" in text else text[3:]
            text = text.rsplit("```", 1)[0]

        return json.loads(text)
    except Exception as e:
        return {"error": str(e)}


def build_cover_letter_docx(
    content: dict,
    profile: dict,
    output_path: str | Path,
) -> Path:
    """Build an ATS-compliant .docx cover letter."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from datetime import date

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    personal = profile.get("personal", {})
    name = f"{personal.get('first_name', '')} {personal.get('last_name', '')}".strip()
    if name:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.font.size = Pt(14)
        run.font.bold = True

    contact_parts = []
    if personal.get("email"):
        contact_parts.append(personal["email"])
    if personal.get("phone"):
        contact_parts.append(personal["phone"])
    if personal.get("location"):
        contact_parts.append(personal["location"])
    if contact_parts:
        doc.add_paragraph(" | ".join(contact_parts)).style.font.size = Pt(10)

    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph()

    doc.add_paragraph(content.get("greeting", "Dear Hiring Team,"))
    doc.add_paragraph()

    if content.get("opening"):
        doc.add_paragraph(content["opening"])

    if content.get("body"):
        doc.add_paragraph(content["body"])

    if content.get("closing"):
        doc.add_paragraph(content["closing"])

    doc.add_paragraph()
    doc.add_paragraph(content.get("sign_off", "Sincerely,"))
    if name:
        doc.add_paragraph(name)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def generate_cover_letter_for_job(
    client: Any,
    job: dict,
    profile: dict,
    model: str = "gpt-4o-mini",
) -> Optional[Path]:
    """Full pipeline: generate content -> build docx -> return path."""
    content = generate_cover_letter_content(
        client=client,
        job_title=job.get("title", ""),
        job_company=job.get("company", ""),
        job_description=job.get("description", ""),
        profile=profile,
        company_research=job.get("company_research", ""),
        model=model,
    )

    if content.get("error"):
        return None

    personal = profile.get("personal", {})
    first = personal.get("first_name", "Cover")
    last = personal.get("last_name", "Letter")
    company_clean = (job.get("company") or "Company").replace(" ", "_").replace("/", "_")
    filename = f"{first}_{last}_CoverLetter_{company_clean}.docx"
    output_path = OUTPUT_DIR / filename

    build_cover_letter_docx(content, profile, output_path)
    return output_path
